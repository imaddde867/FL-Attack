from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

md_path = "results/report/FL_Security_Research_Report.md"
docx_path = "results/report/FL_Security_Research_Report.docx"

with open(md_path, "r", encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()

def add_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text):
    doc.add_paragraph(text)

idx = 0
while idx < len(lines):
    line = lines[idx]
    if line.startswith("# "):
        add_title(line[2:].strip())
        idx += 1
    elif line.startswith("## "):
        add_heading(line[3:].strip(), level=1)
        idx += 1
    elif line.startswith("### "):
        add_heading(line[4:].strip(), level=2)
        idx += 1
    elif line.startswith("#### "):
        add_heading(line[5:].strip(), level=3)
        idx += 1
    elif line.strip().startswith("| ") and "|" in line[2:]:
        table_lines = []
        while idx < len(lines) and lines[idx].strip().startswith("| "):
            table_lines.append(lines[idx])
            idx += 1
        # Remove separator line if present and enough lines exist
        if len(table_lines) > 1 and set(table_lines[1].strip()) <= set("|- "):
            del table_lines[1]
        if not table_lines:
            continue
        headers = [h.strip() for h in table_lines[0].strip().split("|")[1:-1] if h.strip()]
        if not headers:
            continue
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        for row_line in table_lines[1:]:
            if set(row_line.strip()) <= set("|- "):
                continue
            row_cells = [c.strip() for c in row_line.strip().split("|")[1:-1]]
            if len(row_cells) != len(headers):
                continue
            row = table.add_row().cells
            for i, cell in enumerate(row_cells):
                row[i].text = cell
    elif line.strip().startswith("```"):
        code_lines = [line]
        idx += 1
        while idx < len(lines) and not lines[idx].strip().startswith("```"):
            code_lines.append(lines[idx])
            idx += 1
        if idx < len(lines):
            code_lines.append(lines[idx])
            idx += 1
        code_text = "".join(code_lines)
        doc.add_paragraph(code_text, style="Intense Quote")
    elif line.strip() == "---":
        doc.add_page_break()
        idx += 1
    elif line.strip():
        add_paragraph(line.strip())
        idx += 1
    else:
        idx += 1

doc.save(docx_path)
print(f"Word report generated: {docx_path}")
